#! python3
import docx

print('Creating an invitation Word Document with guests.txt...')
doc = docx.Document()

pageNum = 1
with open('guests.txt') as guests:
    for line in guests:
        doc.add_heading('It would be a pleasure to have the company of', 1)
        doc.add_heading(line.rstrip(), 1)
        doc.add_heading('at 11010 Memory Lane at the Evening of', 1)
        doc.add_heading('April 1st', 1)
        doc.add_heading('at 7 o\'clock', 1)
        doc.add_page_break()

doc.save('Invitations.docx')

print('Done!')